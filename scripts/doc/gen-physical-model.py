#!/usr/bin/env python3
"""Harness 通用脚本：根据 JSON 数据生成物理模型设计文档 .docx
用法: python3 gen-physical-model.py <input.json> [output.docx]
模板: ~/.claude/harness/09-templates/系统物理模型设计文档模版.docx (格式权威)
"""

import json, sys, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

HEADERS = ['字段', '字段名称', '类型(长度)', '是否可空', '默认值', '说明']

def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
    shading.append(shd)

def add_table_header(table):
    hdr = table.rows[0]
    for i, text in enumerate(HEADERS):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = '微软雅黑'
        set_cell_shading(cell, 'D9E2F3')

def add_field_row(table, row_idx, field_data):
    row = table.rows[row_idx]
    for i, text in enumerate(field_data):
        cell = row.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(str(text or ''))
        run.font.size = Pt(9)
        run.font.name = '微软雅黑'

def add_section_separator(table, row_idx, label, color='D9E2F3'):
    row = table.rows[row_idx]
    row.cells[0].merge(row.cells[5])
    row.cells[0].merge(row.cells[4])
    row.cells[0].merge(row.cells[3])
    row.cells[0].merge(row.cells[2])
    row.cells[0].merge(row.cells[1])
    cell = row.cells[0]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = '微软雅黑'
    set_cell_shading(cell, color)

def add_index_row(table, row_idx, idx_data):
    row = table.rows[row_idx]
    row.cells[0].merge(row.cells[1])
    row.cells[2].merge(row.cells[3])
    row.cells[4].merge(row.cells[5])
    idx_name, idx_fields, idx_type = idx_data
    for cell, text in [(row.cells[0], idx_name), (row.cells[2], idx_fields), (row.cells[4], idx_type)]:
        cell.text = ''
        run = cell.paragraphs[0].add_run(text)
        run.font.size = Pt(9)
        run.font.name = '微软雅黑'

def add_data_row(table, row_idx, label, value):
    row = table.rows[row_idx]
    row.cells[0].merge(row.cells[2])
    row.cells[3].merge(row.cells[5])
    for cell, text in [(row.cells[0], label), (row.cells[3], value)]:
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        if cell == row.cells[0]:
            run.bold = True
        run.font.size = Pt(9)
        run.font.name = '微软雅黑'

AUDIT_FIELDS = {
    'create_time': ('create_time', '创建时间', 'datetime', '非空', '', 'yyyy-MM-dd hh:mm:ss'),
    'update_time': ('update_time', '更新时间', 'datetime', '非空', '', 'yyyy-MM-dd hh:mm:ss'),
    'create_by': ('create_by', '创建人', 'bigint(20)', '非空', '', ''),
    'update_by': ('update_by', '更新人', 'bigint(20)', '非空', '', ''),
    'delete_flag': ('delete_flag', '删除标识', 'tinyint', '非空', '0', '0-未删除，1-已删除'),
    'tenant_id': ('tenant_id', '租户ID', 'int', '非空', '', '多租户隔离场景租户id'),
}

def build_table(doc, table_def):
    """Build one table from table definition dict"""
    fields = table_def.get('fields', [])
    indexes = table_def.get('indexes', [])
    data_volume = table_def.get('data_volume', '')
    data_retention = table_def.get('data_retention', '')
    has_create_by = table_def.get('has_create_by', False)
    has_update_by = table_def.get('has_update_by', False)
    has_delete_flag = table_def.get('has_delete_flag', True)

    audit_list = ['create_time', 'update_time']
    if has_create_by: audit_list.append('create_by')
    if has_update_by: audit_list.append('update_by')
    if has_delete_flag: audit_list.append('delete_flag')
    audit_list.append('tenant_id')

    n_audit = len(audit_list)
    n_data_pairs = (1 if data_volume else 0) + (1 if data_retention else 0)
    total_rows = 1 + len(fields) + 1 + n_audit + 1 + len(indexes)
    if data_volume:
        total_rows += 2  # separator + value
    if data_retention:
        total_rows += 2

    # Add table heading
    name = table_def.get('name', 'table_name')
    cn_name = table_def.get('cn_name', '表中文名')
    doc.add_heading(f'{name}（{cn_name}）', level=2)

    table = doc.add_table(rows=total_rows, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header(table)

    row_idx = 1

    # Business fields
    for f in fields:
        add_field_row(table, row_idx, f)
        row_idx += 1

    # Other info separator
    add_section_separator(table, row_idx, '其他信息')
    row_idx += 1

    # Audit fields
    for key in audit_list:
        add_field_row(table, row_idx, list(AUDIT_FIELDS[key]))
        row_idx += 1

    # Index separator
    add_section_separator(table, row_idx, '索引')
    row_idx += 1

    for idx in indexes:
        add_index_row(table, row_idx, idx)
        row_idx += 1

    # Data volume + retention
    if data_volume:
        add_section_separator(table, row_idx, '数据增量', 'FFFFFF')
        row_idx += 1
        add_data_row(table, row_idx, '', data_volume)
        row_idx += 1

    if data_retention:
        add_section_separator(table, row_idx, '数据存储时效', 'FFFFFF')
        row_idx += 1
        add_data_row(table, row_idx, '', data_retention)
        row_idx += 1

    doc.add_paragraph()  # spacer

def generate(input_path, output_path=None):
    if not output_path:
        output_path = input_path.replace('.json', '.docx')

    with open(input_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc = Document()
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10)

    # Title
    title = doc.add_heading(data.get('title', '物理模型设计文档'), level=0)
    for run in title.runs:
        run.font.name = '微软雅黑'

    # Metadata
    meta = doc.add_paragraph()
    meta.add_run(f"文档编号: {data.get('doc_id', '')}  |  文档密级: 内部  |  编写人: {data.get('author', '')}  |  编写日期: {data.get('date', '')}").font.size = Pt(9)

    # Version history
    doc.add_heading('版 本 历 史', level=1)
    vh = ['版本/状态', '作者', '变更日期', '备注', '审阅']
    vt = doc.add_table(rows=2, cols=5)
    vt.style = 'Table Grid'
    vt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, text in enumerate(vh):
        c = vt.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(text)
        r.bold = True
        r.font.size = Pt(9)
        r.font.name = '微软雅黑'
        set_cell_shading(c, 'D9E2F3')
    for i, text in enumerate(['V1.0', data.get('author', ''), data.get('date', ''), '初始版本', '']):
        vt.rows[1].cells[i].text = ''
        run = vt.rows[1].cells[i].paragraphs[0].add_run(text)
        run.font.size = Pt(9)
        run.font.name = '微软雅黑'

    # Domains & tables
    for domain in data.get('domains', []):
        domain_name = domain.get('name', '')
        if domain_name:
            doc.add_heading(domain_name, level=1)

        for table_def in domain.get('tables', []):
            build_table(doc, table_def)

    # Appendix
    appendix = data.get('appendix', {})
    if appendix:
        doc.add_heading('附录', level=1)
        for section_title, items in appendix.items():
            doc.add_heading(section_title, level=2)
            for item in items:
                doc.add_paragraph(item, style='List Bullet')

    doc.save(output_path)
    table_count = sum(len(d.get('tables', [])) for d in data.get('domains', []))
    return table_count

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法: python3 gen-physical-model.py <input.json> [output.docx]")
        sys.exit(1)
    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None
    count = generate(input_path, output_path)
    out = output_path or input_path.replace('.json', '.docx')
    print(f"✅ {out}  ({count} tables)")
